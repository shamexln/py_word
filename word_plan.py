import docx
from sortedcontainers import SortedSet

doc = docx.Document('Ataln 软件验证计划.docx')
haids = SortedSet()
for t in doc.tables:
        # j = 0
        for row in t.rows:
            # j = j + 1
            i = 0    
            for cell in row.cells:
                if i == 5:
                    for paragraph in cell.paragraphs:
                        if '通过' in paragraph.text:
                            oldVal = paragraph.text
                            print (oldVal)
                            paragraph.text = '计划中'
                            print (oldVal, paragraph.text)
                            
                i = i + 1
                
doc.save('Ataln 软件验证计划1.docx')               