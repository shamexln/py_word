import docx
from sortedcontainers import SortedSet

doc = docx.Document('CH3.5.5.6 可追溯分析报告.docx')
arr = []
tsr = []
emptyline = 0
found  = 0
for t in doc.tables:
        # j = 0
        for row in t.rows:
            # j = j + 1
            i = 0    
            for cell in row.cells:
                if i == 4:
                    for paragraph in cell.paragraphs:
                        print (paragraph.text)
                        if paragraph.text.startswith('Axxx') :
                            endIndex = paragraph.text.find('(')
                            if endIndex == -1:
                                endIndex = paragraph.text.find('（')
                            if endIndex != -1:
                                arr.append(paragraph.text[:endIndex])
                                print (paragraph.text[:endIndex])
                        elif paragraph.text.startswith('PCS') :
                            found = 1
                    str = ''
                    for tt in arr:
                        str = str + tt + '\n'
                    if len(arr) > 0:
                        cell.text = str
                        arr.clear()
                    elif len(arr) == 0 and found == 1:
                        emptyline = 1
                        found = 0
                i = i + 1
            if emptyline == 1:
                row._element.getparent().remove(row._element)
                emptyline = 0
            # if j == 70:
            #    break;  
doc.save('Traceability_SRS1_MDR_ZH1.docx')               