import docx
from sortedcontainers import SortedSet

doc = docx.Document('Traceability_SRS1_MDR_ZH.docx')
arr = []
tsr = []
for t in doc.tables:
        #j = 0
        for row in t.rows:
            #j = j + 1
            i = 0    
            for cell in row.cells:
                if i == 0:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.startswith('PR_') or paragraph.text.startswith('HA_'):
                            arr.append(paragraph.text)
                            print (paragraph.text)
                    str = ''
                    for tt in arr:
                        str = str + tt + '\n'
                    if len(arr) > 0:
                        cell.text = str
                        arr.clear()
                elif i == 1:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.startswith('TSR_'):
                            tsr.append(paragraph.text)
                            print (paragraph.text)
                    str = ''
                    for tt in tsr:
                        str = str + tt + '\n'
                    if len(tsr) > 0:
                        cell.text = str
                        tsr.clear()
                i = i + 1
            #if j == 70:
               #break;  
doc.save('Traceability_SRS1_MDR_ZH1.docx')               