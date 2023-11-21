import docx
from sortedcontainers import SortedSet

doc = docx.Document('Traceability_SRS1_MDR_ZH_ORG.docx')
haids = SortedSet()
arr = []
for t in doc.tables:
        # j = 0
        for row in t.rows:
            i = 0  
            for cell in row.cells:
                if i == 0:
                    i = i + 1
                    for paragraph in cell.paragraphs:
                        if 'PR_' in paragraph.text:
                            print (i, paragraph.text)
                            ls = paragraph.text.rsplit('_', 1)
                            try:
                                if len(ls) > 1:
                                    print (int(ls[1].strip()))
                                    # haids.add(int(ls[1].strip()))
                                    arr.append(int(ls[1].strip()))
                            except ValueError:
                                print(ls)
                
with open('tracepr_org.txt', 'w') as f:
    # ss = SortedSet(haids)
    ss = arr.copy()
    for id in ss:
        f.write (str(id))
        f.write('\n')
                    