import docx
from sortedcontainers import SortedSet

doc = docx.Document('Traceability_SRS2_IN4k.docx')
haids = SortedSet()
for t in doc.tables:
        # j = 0
        for row in t.rows:
            # j = j + 1    
            for cell in row.cells:
                i = 0
                for paragraph in cell.paragraphs:
                    i = i +1
                    if 'HA_A' in paragraph.text:
                        print (i, paragraph.text)
                        ls = paragraph.text.rsplit('_', 1)
                        try:
                            if len(ls) > 1:
                                print (int(ls[1].strip()))
                                haids.add(int(ls[1].strip()))
                        except ValueError:
                            print(ls)
            # if j == 60:
            #     break;
                
                
with open('trace_in4k.txt', 'w') as f:
    ss = SortedSet(haids)
    for id in ss:
        f.write (str(id))
        f.write('\n')
                    