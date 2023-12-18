import docx
from sortedcontainers import SortedSet

doc = docx.Document('CH3.5.5.6 可追溯分析报告.docx')
haids = SortedSet()
for t in doc.tables:
        # j = 0
        for row in t.rows:
            # j = j + 1    
            for cell in row.cells:
                i = 0
                for paragraph in cell.paragraphs:
                    if i == 0:
                        if paragraph.text.startswith('HA_'):
                            print (i, paragraph.text)
                            lsarr = paragraph.text.split('\n')
                            count = len(lsarr)
                            try:
                                for i in  range(0, count):
                                    print ((lsarr[i].strip()))
                                    subls = lsarr[i].strip().rsplit('_', 1)
                                    try:
                                        if len(subls) > 1: 
                                            print (int(subls[1].strip()))
                                            haids.add(int(subls[1].strip()))
                                    except ValueError:
                                        print(subls)
                                    i = i + 1
                            except ValueError:
                                print(lsarr)
                    i = i + 1
            # if j == 60:
            #     break;
                
                
with open('CH3.5.5.6.ID.txt', 'w') as f:
    ss = SortedSet(haids)
    for id in ss:
        f.write (str(id))
        f.write('\n')
                    